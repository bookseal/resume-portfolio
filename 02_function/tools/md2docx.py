"""마크다운 → docx 변환기 (이력서·자기소개서 공용).

기존 build_resume_docx.py는 이력서 '내용'이 코드에 하드코딩돼 있어서 회사가 바뀔 때마다
스크립트를 통째로 복사해야 했다. 이 스크립트는 내용을 md에서 읽어오므로 회사마다 md만
새로 쓰면 된다. 스타일(폰트·여백·1페이지 압축)은 여기 한 곳에서만 관리한다.

Google Docs 업로드 호환을 위해 표/텍스트박스/다단 없이 내장 스타일만 사용한다.

사용법:
    python function/tools/md2docx.py <입력.md> [출력.docx]
    python function/tools/md2docx.py output/LGUplus_DAX_DevOpsSRE/이력서_이기찬_LGUplus_DAX_20260713_1530.md

출력 경로를 생략하면 입력 md와 같은 폴더에 같은 이름의 .docx를 만든다.

지원하는 마크다운:
    # 제목        → 가운데 정렬 이름 (문서 최상단 1회)
    ## 섹션       → Heading 1
    ### 하위섹션   → Heading 2
    - 항목        → List Bullet
    **굵게**      → bold run
    *기울임*      → 메타 라인 (날짜·소속 등)
    ---          → 무시 (구분선)
    > 인용        → 본문으로 처리
"""
import re
import sys
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

LATIN_FONT = "Calibri"
KO_FONT = "맑은 고딕"


def set_fonts(doc):
    """기본 스타일 폰트 지정. 한글은 eastAsia 속성을 직접 박아야 적용된다."""
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet"):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = LATIN_FONT
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), KO_FONT)
        rfonts.set(qn("w:ascii"), LATIN_FONT)
        rfonts.set(qn("w:hAnsi"), LATIN_FONT)

    doc.styles["Normal"].font.size = Pt(10)

    for name, size, color in (
        ("Heading 1", 13, RGBColor(0x1F, 0x1F, 0x1F)),
        ("Heading 2", 11, RGBColor(0x1F, 0x1F, 0x1F)),
    ):
        try:
            style = doc.styles[name]
        except KeyError:
            continue
        style.font.size = Pt(size)
        style.font.bold = True
        style.font.color.rgb = color


def set_margins(doc, inches=0.7):
    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(inches)
        section.right_margin = Inches(inches)


def add_runs(paragraph, text):
    """**굵게** 마크업만 해석해 run으로 나눈다."""
    for i, chunk in enumerate(text.split("**")):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        run.bold = i % 2 == 1


def add_centered(doc, text, size, bold=False, color=None, space_after=2):
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(space_after)
    run = p.add_run(text)
    run.bold = bold
    run.font.size = Pt(size)
    if color:
        run.font.color.rgb = color
    return p


def strip_md(text):
    """docx로 옮길 때 의미 없는 마크업(링크·인라인 코드)을 벗긴다."""
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)  # [텍스트](url) → 텍스트
    text = text.replace("`", "")
    return text.strip()


def build_header(doc, lines):
    """문서 최상단의 이름·헤드라인·연락처 블록을 가운데 정렬로 만든다.

    `# 이름` 이후 첫 `---` 까지를 헤더로 본다.
    - 첫 줄(# ...)          → 이름 (20pt bold)
    - **굵은** 줄           → 헤드라인 (10.5pt bold)
    - 나머지 줄             → 연락처·링크 (9pt 회색), 여러 줄이면 줄바꿈으로 합침
    """
    name = lines[0].lstrip("# ").strip()
    add_centered(doc, name, size=20, bold=True)

    contact = []
    for line in lines[1:]:
        line = strip_md(line)
        if not line:
            continue
        if line.startswith("**") and line.endswith("**"):
            add_centered(doc, line.strip("*"), size=10.5, bold=True)
        else:
            contact.append(line)

    if contact:
        add_centered(
            doc,
            "\n".join(contact),
            size=9,
            color=RGBColor(0x44, 0x44, 0x44),
            space_after=8,
        )


def convert(md_path: Path, out_path: Path):
    raw = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document()
    set_fonts(doc)
    set_margins(doc)

    # --- 헤더 블록: 첫 '---' 전까지 ---
    try:
        split_at = raw.index("---")
    except ValueError:
        split_at = 0
    if split_at:
        build_header(doc, [l for l in raw[:split_at] if l.strip()])
        body_lines = raw[split_at + 1 :]
    else:
        body_lines = raw

    # --- 본문 ---
    for line in body_lines:
        line = line.rstrip()
        stripped = line.strip()

        if not stripped or stripped == "---":
            continue

        if stripped.startswith("<!--"):  # TODO 주석 등은 문서에 넣지 않는다
            continue

        if stripped.startswith("### "):
            doc.add_heading(strip_md(stripped[4:]), level=2)

        elif stripped.startswith("## "):
            doc.add_heading(strip_md(stripped[3:]), level=1)

        elif stripped.startswith("- "):
            p = doc.add_paragraph(style="List Bullet")
            p.paragraph_format.space_after = Pt(2)
            add_runs(p, strip_md(stripped[2:]))

        elif stripped.startswith("*") and stripped.endswith("*") and "**" not in stripped:
            # *2026.05 – 현재* 같은 메타 라인
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(3)
            run = p.add_run(strip_md(stripped.strip("*")))
            run.italic = True
            run.font.size = Pt(9)
            run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

        else:
            p = doc.add_paragraph()
            p.paragraph_format.space_after = Pt(4)
            add_runs(p, strip_md(stripped.lstrip("> ")))

    doc.save(out_path)
    return out_path


def main():
    if len(sys.argv) < 2:
        print(__doc__)
        sys.exit(1)

    md_path = Path(sys.argv[1])
    if not md_path.exists():
        sys.exit(f"입력 파일이 없습니다: {md_path}")

    out_path = Path(sys.argv[2]) if len(sys.argv) > 2 else md_path.with_suffix(".docx")

    convert(md_path, out_path)
    print(f"생성: {out_path}")


if __name__ == "__main__":
    main()
