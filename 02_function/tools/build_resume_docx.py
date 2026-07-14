"""이력서 md → docx 변환. Google Docs 업로드 호환을 위해 내장 스타일만 사용한다.

표/텍스트박스/다단 없이 Heading·List Bullet·Normal 세 가지로만 구성.
"""
from datetime import datetime
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.shared import Pt, RGBColor, Inches

BASE = Path(__file__).parent
STAMP = datetime.now().strftime("%Y%m%d_%H%M")
OUT = BASE / f"이력서_이기찬_AISolutionEngineer_218421_{STAMP}.docx"

LATIN_FONT = "Calibri"
KO_FONT = "맑은 고딕"


def set_fonts(doc):
    """기본 스타일 폰트 지정. 한글은 eastAsia 속성을 직접 박아야 적용된다."""
    for style_name in ("Normal", "Heading 1", "Heading 2", "Heading 3", "List Bullet"):
        try:
            style = doc.styles[style_name]
        except KeyError:
            continue
        style.font.name = LATIN_FONT
        rpr = style.element.get_or_add_rPr()
        rfonts = rpr.get_or_add_rFonts()
        rfonts.set(qn("w:eastAsia"), KO_FONT)
        rfonts.set(qn("w:ascii"), LATIN_FONT)
        rfonts.set(qn("w:hAnsi"), LATIN_FONT)

    normal = doc.styles["Normal"]
    normal.font.size = Pt(10)
    normal.paragraph_format.space_after = Pt(2)
    normal.paragraph_format.line_spacing = 1.15

    for name, size in (("Heading 1", 13), ("Heading 2", 11), ("Heading 3", 10)):
        st = doc.styles[name]
        st.font.size = Pt(size)
        st.font.bold = True
        st.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
        st.paragraph_format.space_before = Pt(10 if name == "Heading 1" else 6)
        st.paragraph_format.space_after = Pt(3)


def add_runs(paragraph, text):
    """**굵게** 마크업만 해석해 run으로 나눈다."""
    for i, chunk in enumerate(text.split("**")):
        if not chunk:
            continue
        run = paragraph.add_run(chunk)
        run.bold = i % 2 == 1


def main():
    doc = Document()
    set_fonts(doc)

    for section in doc.sections:
        section.top_margin = Inches(0.6)
        section.bottom_margin = Inches(0.6)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # 헤더
    name = doc.add_paragraph()
    name.alignment = WD_ALIGN_PARAGRAPH.CENTER
    name.paragraph_format.space_after = Pt(2)
    run = name.add_run("이기찬 (Gichan Lee)")
    run.bold = True
    run.font.size = Pt(20)

    tagline = doc.add_paragraph()
    tagline.alignment = WD_ALIGN_PARAGRAPH.CENTER
    tagline.paragraph_format.space_after = Pt(2)
    run = tagline.add_run(
        "AI Solution Engineer — 고객 인프라에 AI를 배포하고, 파트너가 스스로 운영하게 만드는 현직 FDE"
    )
    run.bold = True
    run.font.size = Pt(10.5)

    contact = doc.add_paragraph()
    contact.alignment = WD_ALIGN_PARAGRAPH.CENTER
    contact.paragraph_format.space_after = Pt(8)
    run = contact.add_run(
        "010-6666-2824 | gichanlee@icloud.com | 서울시 강남구 자곡동\n"
        "github.com/bookseal | bit-habit.com | wiki.bit-habit.com | linkedin.com/in/bookseal"
    )
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x44, 0x44, 0x44)

    def h1(text):
        doc.add_heading(text, level=1)

    def h2(text):
        doc.add_heading(text, level=2)

    def meta(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(3)
        run = p.add_run(text)
        run.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

    def bullet(text):
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        add_runs(p, text)

    def body(text):
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_runs(p, text)

    # 요약
    h1("요약")
    bullet(
        "**온프레미스 k8s 배포·운영·트러블슈팅** — OCI 단일 k3s 클러스터에 14개 서비스를 직접 구성해 2년째 운영 중. "
        "ArgoCD GitOps 기반 CI/CD, Prometheus 모니터링 스택 구축 및 장애 대응."
    )
    bullet(
        "**파트너·비개발 이해관계자 기술 지원과 교육** — MS M365 글로벌 기술지원에서 20명 중 전 지표 1위"
        "(해결률 92.7%, 불만족 0%). 팀 멘토십(MS-900) 운영. 42 Seoul 3년 스터디 운영."
    )
    bullet(
        "**복잡한 엔지니어링을 문서로 자산화** — 클라우드/AI 기술 가이드북(wiki.bit-habit.com)을 직접 집필·공개. "
        "비전공자 온보딩과 정부 지침서 LLM 워크플로우화(10h→2h)를 현직에서 수행."
    )

    # 경력
    h1("경력")

    h2("한국경영분석연구원 | AI Technical Product Manager / Lead AI FDE")
    meta("2026.05 – 현재")
    bullet(
        "**AI Agent 기반 업무 자동화 파이프라인 구축** — Claude Code·Codex를 사내 개발 프로세스에 도입, "
        "GitHub Issue 생성부터 PR 작성·코드 리뷰까지 자동화. Teams 실시간 자막(STT)을 NotebookLM과 연동해 "
        "회의록 요약·검색·문서화를 자동화."
    )
    bullet(
        "**정부 지침서 기반 LLM 워크플로우 설계** — 소프트웨어 대가 산정 가이드 분석 업무를 10시간 → 2시간으로 단축. "
        "비전공자 팀원이 스스로 돌릴 수 있도록 절차를 문서화·온보딩."
    )
    bullet(
        "**사내 온프레미스 서버 및 배포 자동화 운영** — Ubuntu 서버에 Docker/k3s 환경을 구축하고 GitHub PR 병합 시 "
        "자동 배포되는 CI/CD 파이프라인 운영. 팀원 장비 성능 한계는 OCI·Codespaces 기반 웹 IDE로 해결."
    )
    bullet(
        "**요구사항 구조화 및 제품 피드백 루프** — 경영진의 비즈니스 요구를 기술 과제로 번역하고 스코프를 관리. "
        "사내 HR·자격증 추천 시스템(Quali-fit) 기획부터 배포까지 리드."
    )
    bullet(
        "**반복 병목의 근본 해결** — 사내 요구사항·피드백 수집을 위해 대시보드를 하드코딩하는 대신 검증된 "
        "오픈소스 피드백 포털(self-hosted)을 발굴·서버 세팅해 개발 리소스를 대폭 절감. 매년 갱신되는 대가 산정 "
        "가이드는 크롤링·스프레드시트 동기화 파이프라인으로 수작업 제거."
    )
    bullet(
        "**민감정보 처리 자동화** — 화면 캡처 이미지 내 개인정보를 Vision 모델이 자동 인식·마스킹(Redaction)하는 "
        "기능을 업무 워크플로우에 적용."
    )
    bullet(
        "**팀 기술 교육 및 온보딩** — 신규 합류 주니어에게 Git 브랜치 전략·클라우드 개발 환경 세팅을 화면공유로 "
        "직접 교육. 비개발 경영진 대상 시연은 롤플레잉 방식의 모의 발표로 훈련시켜 팀의 전달 역량을 강화."
    )
    bullet(
        "**원격·시차 환경의 비동기 협업 체계 정립** — 12시간 시차(시애틀) 하에서 GitHub PR 중심 비동기 병합 "
        "프로세스를 설계해 개발이 끊기지 않는 환경 구성."
    )

    h2("Bithabit | AI Architect / Infrastructure Owner (Founder)")
    meta("2024.08 – 현재")
    bullet(
        "**단일 k3s 클러스터에 14개 서비스 운영 (bit-habit-infra)** — Oracle Cloud 상에 클러스터를 직접 구성하고 "
        "네트워크·스토리지·인그레스 설계. AWS → OCI 이관으로 인프라 비용을 $0로 유지."
    )
    bullet(
        "**ArgoCD 기반 GitOps 배포 자동화 및 관찰가능성 확보** — 선언형 매니페스트로 배포 파이프라인을 표준화하고, "
        "Prometheus 모니터링 스택을 구성해 서비스 상태를 상시 관측."
    )
    bullet(
        "**웹 인증·보안 직접 구현** — Email OTP + JWT 기반 인증/인가 플로우를 설계·구현하고 2년 이상 "
        "실사용 트래픽에서 운영."
    )
    bullet(
        "**Bithabit 습관 형성 서비스** — IITP 해커톤 2위, 2년 이상 실사용 서비스로 운영 중. 서버 비용 절감을 위해 "
        "타임랩스 GIF 생성을 클라이언트로 이전($0)."
    )

    h2("Concentrix Korea (Microsoft M365) | Technical Success Advisor")
    meta("2025.08 – 2025.11")
    bullet(
        "**글로벌 고객 기술지원 전 지표 팀 1위 (20명 중)** — 문제 해결률 92.7%, 불만족(DSAT) 0%, 고객 별점 42건. "
        "Microsoft Global Tester로 선발."
    )
    bullet("**Fireteam 멘토십 운영 (MS-900)** — 신규 인입 엔지니어 대상 기술 교육 세션을 설계·진행.")
    bullet(
        "라이선스·인증·권한(OAuth·테넌트 관리) 및 네트워크 이슈를 영어로 진단하고, 재현 절차와 해결 과정을 "
        "문서화해 팀 자산으로 축적."
    )

    h2("FPT Software Korea | Embedded Support Engineer")
    meta("2024.10 – 2025.07")
    bullet(
        "현대오토에버 차량 보안 모듈 진단 및 리눅스 환경 트러블슈팅 지원. 이슈 재현 절차와 로그 분석 과정을 "
        "정리해 공유하며, 담당 범위를 넘어 팀원들의 환경 이슈 해결을 상시 지원."
    )

    h2("가마솥 화상영어 (베트남 하노이) | Founder & Operator")
    meta("2018 – 2022 (약 4년)")
    bullet(
        "**다국적 파트너 운영** — 한국 학생과 필리핀 교사를 연결하는 완전 비대면 교육 사업을 4년간 운영. "
        "파트너(교사) 채용·교수법 코칭·품질 관리를 직접 수행하고, 교사 관리 도구와 Unity 학습 게임을 자체 개발."
    )
    bullet(
        "**사용자 피드백 루프 운영** — 카카오톡 채널로 학생·학부모·교사와 상시 소통하며 현장의 요구를 수집·구조화해 "
        "커리큘럼과 운영 방식에 반영."
    )

    # 프로젝트
    h1("주요 프로젝트")

    h2("BookToss — booktoss.bit-habit.com")
    body(
        "LangGraph + Browser-use/Playwright로 공개 API가 없는 서울 25개 구립도서관을 통합 검색하는 AI 에이전트. "
        "도서 탐색 시간 5분 → 1분(80%↓). 현재 Upstage Solar API로 추론 엔진 전면 재구축 중."
    )

    h2("Quali-fit — quali-fit.bit-habit.com")
    body(
        "사내 HR·자격증 추천 시스템. 엑셀/VBA 수작업으로 흩어져 있던 직원·학력·자격증·업무분류 데이터를 SQLite로 "
        "통합하고 Streamlit 대시보드로 전환. 자격증 만료 경고, 인쇄 최적화 출력 등 최종 사용자(경영진)의 실제 "
        "업무 방식에 맞춰 설계하고 배포까지 완료."
    )

    h2("Sentinel — sentinel.bit-habit.com")
    body(
        "메가존클라우드 해커톤 1위. 실시간 음성 볼륨 감지를 로컬 NumPy 연산으로 처리해 추론 비용 $0·지연 10ms 미만 달성."
    )

    h2("Cloud Architect's Playbook — wiki.bit-habit.com")
    body(
        "클라우드·M365·AI Agent Engineering·아키텍트를 위한 시스템 수학을 다루는 기술 가이드북. "
        "복잡한 개념을 누구나 따라올 수 있는 순서로 재구성해 공개."
    )

    h2("Seoul APT Prediction / Viz Platform — seoul-apt.bit-habit.com, viz.bit-habit.com")
    body(
        "머신러닝을 휴리스틱 → 선형회귀 → PCA → AutoML까지 10단계로, 선형대수를 벡터 → SVD → 임베딩까지 "
        "시각화로 체험하게 만든 학습 플랫폼 (Streamlit, scikit-learn, Manim)."
    )

    # 기술 스택
    h1("기술 스택")
    bullet(
        "**Infra & Ops**: Kubernetes(k3s), Docker, ArgoCD(GitOps), Prometheus, Linux(Ubuntu), Oracle Cloud, "
        "GitHub Actions CI/CD, Nginx"
    )
    bullet("**Language & Scripting**: Python, Bash, TypeScript/JavaScript, C")
    bullet(
        "**AI & LLM**: Solar API, LangGraph, RAG, Claude Code / Codex 기반 에이전트 워크플로우, "
        "Browser-use, Playwright"
    )
    bullet("**Web & Data**: HTTP/REST, JSON, JWT·OAuth 인증, Next.js, Streamlit, Flutter, SQLite/PostgreSQL")
    bullet("**Communication**: 비즈니스 영어 (글로벌 기술지원·해외 파트너 운영), 기술 문서 작성")

    # 학력
    h1("학력 & 활동")
    bullet(
        "**École 42 Seoul** — Computer Science (2022.10 – 2024.10). UNIX 시스템·네트워크 프로그래밍. "
        "3년간 동료 스터디 운영, 동료 1명의 Upstage AI Research Engineer 인턴 합격을 지원."
    )
    bullet("**부경대학교** — 산업공학 학사 (2007.03 – 2014.02)")
    bullet("**Fast Campus × Upstage AI Lab 수료** — Solar, Document Parse 실습")
    bullet("개발자 모임·해커톤 자원봉사 다수. 국내 출장 가능.")

    doc.save(OUT)
    print(f"saved: {OUT}")


if __name__ == "__main__":
    main()
