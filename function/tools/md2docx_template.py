"""마크다운 → docx. **기존 제출본 docx 를 서식 템플릿으로 재활용한다.**

md2docx.py 는 docx 를 맨바닥에서 만든다(Google Docs 호환 우선). 이 스크립트는 반대로,
워드에서 손봐 둔 제출본 docx 를 열어 **본문만 비우고** md 내용으로 다시 채운다.

왜 이 방식인가:
    워드 서식은 styles.xml(폰트·크기), numbering.xml(글머리표 정의), sectPr(여백) 세 곳에
    흩어져 있다. 코드로 재현하면 폰트·글머리표 모양·여백이 미묘하게 틀어진다.
    원본을 템플릿으로 열면 그 정의가 통째로 딸려온다 — 정확하고, 코드도 짧다.

**서식 매핑을 하드코딩하지 않는다.** 템플릿마다 어떤 스타일이 섹션이고 어떤 게 불릿인지 다르다
(표준서식: 섹션=Heading 1 / 불릿=List Bullet 9pt · 구서식: 섹션=Heading 2 17pt / 불릿='normal' 12pt).
그래서 sniff() 가 템플릿을 읽어 매핑을 스스로 알아낸다:

    · 불릿   = numPr(글머리표)이 붙은 단락의 스타일·크기
    · 회사   = 탭(\\t)이 들어간 제목 단락  ← 두 템플릿 모두 "회사 | 직무 ⟶ 날짜" 에만 탭을 쓴다
    · 섹션   = 탭이 없는 제목 단락
    · 헤더   = 첫 제목 이전의 단락들 (이름 / 헤드라인 / 연락처)

템플릿을 새로 갈아도 코드는 그대로다.

사용법:
    python3 function/tools/md2docx_template.py <입력.md> [출력.docx] [--template <서식.docx>]

기본 템플릿: function/templates/이력서_표준서식.docx
    (= 2026-07-13 Upstage 218421 제출본. 워드에서 1페이지에 맞춰 다듬은 판)
"""
import re
import sys
from copy import deepcopy
from pathlib import Path

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT
from docx.oxml.ns import qn
from docx.shared import Pt, Inches

REPO = Path(__file__).resolve().parents[2]
DEFAULT_TEMPLATE = REPO / "function/templates/이력서_표준서식.docx"

DATE_RE = re.compile(r"^\*\*(.+?)\*\*$")  # **2026.05 – 현재**


# ── 템플릿 서식 파악 ────────────────────────────────────────────────

def _size(par):
    """단락 첫 run 의 크기(pt). None 이면 스타일 기본값을 그대로 쓴다는 뜻."""
    for r in par.runs:
        if r.text.strip():
            return r.font.size.pt if r.font.size else None
    return None


def _has_numPr(el):
    if el is None:
        return False
    pPr = el.find(qn("w:pPr"))
    return pPr is not None and pPr.find(qn("w:numPr")) is not None


def _is_bullet(par):
    """글머리표는 두 층위 중 하나에 있다 — 둘 다 봐야 한다.

    · 단락 직접 서식 (구서식): 단락마다 numPr 이 박혀 있다
    · 스타일 정의   (표준서식): 'List Bullet' 스타일이 numPr 을 들고 있고 단락은 비어 있다

    한쪽만 검사하면 다른 템플릿에서 조용히 폴백된다(에러가 안 난다).
    """
    return _has_numPr(par._p) or _has_numPr(par.style.element)


def sniff(doc):
    """템플릿에서 서식 매핑을 읽어낸다."""
    prof = {
        "header": [],        # [(정렬, 크기)] — 이름/제목, 헤드라인, 연락처…
        "section": None,     # (스타일명, 크기) — ## 섹션
        "company": None,     # (스타일명, 크기) — ### 회사 | 직무  (탭으로 판별)
        "subsection": None,  # (스타일명, 크기) — #### 프로젝트 제목 (탭 없는 두 번째 제목 층위)
        "bullet": None,      # (스타일명, 크기) — - 항목
        "bullet_pPr": None,  # 글머리표+들여쓰기 XML (스타일만으론 글머리표가 안 나온다)
        "body": None,        # (스타일명, 크기) — 산문 문단 (자기소개서)
    }
    seen_heading = False

    for par in doc.paragraphs:
        if not par.text.strip():
            continue
        style = par.style.name
        is_heading = style.startswith("Heading")

        if _is_bullet(par):
            if prof["bullet"] is None:
                prof["bullet"] = (style, _size(par))
                # 단락에 numPr 이 직접 박힌 템플릿에서만 pPr 을 떠 온다.
                # 스타일이 글머리표를 들고 있는 템플릿은 스타일 이름만으로 충분하다.
                pPr = par._p.find(qn("w:pPr"))
                prof["bullet_pPr"] = deepcopy(pPr) if _has_numPr(par._p) else None
            continue

        if is_heading:
            seen_heading = True
            if "\t" in par.text:                      # 탭 = 날짜를 우측으로 미는 회사 줄
                if prof["company"] is None:
                    prof["company"] = (style, _size(par))
            elif prof["section"] is None:             # 탭 없는 첫 제목 = 섹션 (요약·경력·…)
                prof["section"] = (style, _size(par))
            elif style != prof["section"][0] and prof["subsection"] is None:
                prof["subsection"] = (style, _size(par))   # 섹션과 다른 탭 없는 제목 = 프로젝트 제목
            continue

        if not seen_heading:
            prof["header"].append((par.alignment, _size(par)))   # 첫 제목 이전 = 헤더 블록
        elif prof["body"] is None:
            prof["body"] = (style, _size(par))                   # 제목 이후의 산문 문단

    # 표준서식은 이름이 Heading 이 아니라 가운데정렬 Normal 이다 → 헤더가 비면 폴백
    if not prof["header"]:
        prof["header"] = [(WD_ALIGN_PARAGRAPH.CENTER, 17), (WD_ALIGN_PARAGRAPH.CENTER, None)]
    if prof["section"] is None:
        prof["section"] = ("Heading 1", None)
    if prof["company"] is None:
        prof["company"] = ("Heading 2", None)
    if prof["bullet"] is None:
        prof["bullet"] = ("List Bullet", None)
    if prof["subsection"] is None:
        prof["subsection"] = prof["company"]   # 제목 층위가 2개뿐인 템플릿 → 회사 서식을 따른다
    if prof["body"] is None:
        # 이력서처럼 본문이 전부 불릿인 템플릿 → 산문도 불릿 서식을 따른다
        prof["body"] = prof["bullet"]
    return prof


def usable_width_inches(doc):
    """좌우 여백을 뺀 글 영역 폭. 날짜용 우측 탭 정지를 여기에 건다."""
    xml = doc.sections[0]._sectPr.xml
    page = re.search(r'w:w="([\d.]+)"', xml)
    left = re.search(r'w:left="([\d.]+)"', xml)
    right = re.search(r'w:right="([\d.]+)"', xml)
    if not (page and left and right):
        return 6.5
    return (float(page.group(1)) - float(left.group(1)) - float(right.group(1))) / 1440.0


# ── 본문 생성 ──────────────────────────────────────────────────────

def clear_body(doc):
    """본문 단락을 모두 제거한다. sectPr(여백)은 body 의 마지막 자식이라 남긴다."""
    body = doc.element.body
    for child in list(body):
        if not child.tag.endswith("}sectPr"):
            body.remove(child)


def strip_md(text):
    text = re.sub(r"\[([^\]]+)\]\([^)]+\)", r"\1", text)     # [텍스트](url) → 텍스트
    text = re.sub(r"(?<!\*)\*([^*]+)\*(?!\*)", r"\1", text)  # *이탤릭* → 이탤릭 (굵게는 보존)
    return text.replace("`", "").strip()


def add_runs(par, text, size):
    """**굵게** 마크업을 run 으로 나눈다. size=None 이면 스타일 기본 크기를 따른다."""
    for i, chunk in enumerate(text.split("**")):
        if not chunk:
            continue
        run = par.add_run(chunk)
        run.bold = i % 2 == 1
        if size:
            run.font.size = Pt(size)


def add_bullet(doc, prof, text):
    style, size = prof["bullet"]
    par = doc.add_paragraph(style=style)
    if prof["bullet_pPr"] is not None:
        old = par._p.find(qn("w:pPr"))
        if old is not None:
            par._p.remove(old)
        par._p.insert(0, deepcopy(prof["bullet_pPr"]))
    add_runs(par, text, size)
    return par


def add_company(doc, prof, title, date, right_edge):
    """회사 | 직무 ······· 날짜 — 우측 정렬 탭 정지로 날짜를 오른쪽 끝에 붙인다.

    템플릿 원본은 리터럴 탭을 여러 개 박아 밀어냈다. 회사명 길이가 바뀌면 줄이 밀리므로
    여기서는 진짜 우측 탭 정지를 건다 — 길이와 무관하게 항상 오른쪽 끝에 정렬된다.
    """
    style, size = prof["company"]
    par = doc.add_paragraph(style=style)
    if date:
        par.paragraph_format.tab_stops.add_tab_stop(Inches(right_edge), WD_TAB_ALIGNMENT.RIGHT)
    add_runs(par, title, size)
    if date:
        run = par.add_run("\t" + date)
        run.bold = False
        if size:
            run.font.size = Pt(max(size - 1.5, 8))
    return par


def convert(md_path: Path, out_path: Path, template: Path):
    raw = md_path.read_text(encoding="utf-8").splitlines()

    doc = Document(str(template))
    prof = sniff(doc)              # 본문을 비우기 '전'에 서식을 읽는다
    right_edge = usable_width_inches(doc)
    clear_body(doc)

    # 첫 '---' 전까지가 헤더 블록
    try:
        split_at = raw.index("---")
    except ValueError:
        split_at = 0
    header_lines = [l for l in raw[:split_at] if l.strip()] if split_at else []
    body = raw[split_at + 1 :] if split_at else raw

    for i, line in enumerate(header_lines):
        text = strip_md(line).lstrip("# ").strip()
        align, size = prof["header"][min(i, len(prof["header"]) - 1)]
        par = doc.add_paragraph()
        par.alignment = align if align is not None else WD_ALIGN_PARAGRAPH.CENTER
        add_runs(par, text, size)

    skip_next = False
    for idx, line in enumerate(body):
        if skip_next:
            skip_next = False
            continue

        s = line.strip()
        if not s or s == "---" or s.startswith("<!--"):
            continue

        # '####' 를 '###' 보다 먼저 본다 — startswith("### ") 는 '#### ' 도 잡지 못하지만,
        # 순서를 바꾸면 '#### X' 가 '# ' 계열 검사에 먼저 걸릴 수 있다.
        if s.startswith("#### "):
            style, size = prof["subsection"]
            add_runs(doc.add_paragraph(style=style), strip_md(s[5:]), size)

        elif s.startswith("### "):
            title = strip_md(s[4:])
            date = ""
            nxt = body[idx + 1].strip() if idx + 1 < len(body) else ""
            m = DATE_RE.match(nxt)
            if m:                        # 다음 줄이 **날짜** 면 같은 줄 우측에 붙인다
                date = strip_md(m.group(1))
                skip_next = True
            add_company(doc, prof, title, date, right_edge)

        elif s.startswith("## "):
            style, size = prof["section"]
            add_runs(doc.add_paragraph(style=style), strip_md(s[3:]), size)

        elif s.startswith("- "):
            add_bullet(doc, prof, strip_md(s[2:]))

        elif prof["body"] == prof["bullet"]:
            # 이력서: 본문이 전부 불릿인 템플릿 — 산문 줄도 불릿으로 붙인다
            add_bullet(doc, prof, strip_md(s.lstrip("> ")))

        else:
            # 자기소개서: 산문 문단은 글머리표 없이 본문 스타일로
            style, size = prof["body"]
            add_runs(doc.add_paragraph(style=style), strip_md(s.lstrip("> ")), size)

    doc.save(out_path)
    return prof


def main():
    args = list(sys.argv[1:])
    template = DEFAULT_TEMPLATE
    if "--template" in args:
        i = args.index("--template")
        template = Path(args[i + 1])
        del args[i : i + 2]

    if not args:
        print(__doc__)
        sys.exit(1)

    md_path = Path(args[0])
    if not md_path.exists():
        sys.exit(f"입력 파일이 없습니다: {md_path}")
    if not template.exists():
        sys.exit(f"템플릿 docx 가 없습니다: {template}")

    out_path = Path(args[1]) if len(args) > 1 else md_path.with_suffix(".docx")
    prof = convert(md_path, out_path, template)

    print(f"생성: {out_path}")
    print(f"서식: {template.name}")
    print(f"  섹션={prof['section']}  회사={prof['company']}  불릿={prof['bullet']}")


if __name__ == "__main__":
    main()
