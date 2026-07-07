# -*- coding: utf-8 -*-
"""원본 한국어 docx → Cognition FDE 지원용 영어 이력서.
Upstage/교육·DevRel 색채 제거, FDE(비즈니스→배포, 에이전트 시스템) 중심으로 압축."""
import docx
from docx.shared import Twips
from docx.enum.text import WD_TAB_ALIGNMENT

SRC = '이기찬_이력서_빌더_업스테이지 (1).docx'
DST = 'Gichan_Lee_Resume_FDE_Cognition.docx'

R = {
    (0, 0): 'Gichan Lee',
    (1, 0): 'Forward Deployed Engineer — shipping AI agent systems from business intent to production',
    (2, 0): '+82 10-6666-2824 | ',
    (2, 1): ' | Jagok-dong, Gangnam-gu, Seoul, South Korea',
    (5, 0): 'Current Lead AI FDE (AI Agent Lead)',
    (5, 1): ": lead a 3-person team, translating executives' business needs into deployed technology and setting the development direction",
    (6, 0): 'AI agent systems in production',
    (6, 1): ': human-in-the-loop guardrails, meeting-notes → agent → GitHub pipeline, running reliably across a 12-hour US–Korea time gap',
    (7, 0): '#1 in client communication & troubleshooting',
    (7, 1): ': ranked #1 on the team across all metrics (SRR 92.7%, 0% dissatisfaction) in MS M365 enterprise technical support',
    (9, 0): 'Experience',
    (10, 0): 'Korea Institute of Business Analysis (KIBA) | Lead AI FDE',
    (10, 1): '\t',
    (10, 2): '2026.05 - Present',
    (11, 0): '3-Role Orchestration:',
    (11, 1): ' Designed and operate a collaboration loop connecting non-developer clients (intent), AI agents (execution), and verifying developers — agent work made clear enough for non-developers to drive, developer judgment encoded as explicit policy.',
    (12, 0): 'Human-in-the-Loop Guardrails: ',
    (12, 1): 'Agent actions branched by type — reads run automatically, board writes need per-item human approval, deletions/permissions/code changes are prohibited. Meeting notes kept private, PII masked before commits.',
    (13, 0): 'Claude Code + GitHub Pipeline: ',
    (13, 1): 'Asynchronous flow from meeting notes → AI agent → board (Issues/PRs), stable across a 12-hour US–Korea time difference. Migrated Plane → GitHub Projects on a build-vs-buy call.',
    (14, 0): 'Product from Planning to Deployment: ',
    (14, 1): "Built 'Quali-fit', a certification-based staffing recommendation system scoring \"who to assign, and why\". An LLM workflow for government guideline analysis cut a 10-hour task to 2 hours.",
    (15, 0): 'Technical Leadership: ',
    (15, 1): 'Onboarded junior engineers to cloud and Git; ran mock demos so non-technical executives could validate the product before release.',
    (16, 1): '\t',
    (17, 0): 'Achieved a 5.0 average rating, 0% dissatisfaction, and a 92.7% Service Request Resolution (SRR) rate among 20+ engineers (#1 across all metrics)',
    (18, 0): "Founded the internal 'Fireteam' study group (MS-900 certification) and served as a global product tester",
    (19, 1): '\t',
    (19, 2): '2024.09 – 2025.07',
    (20, 0): 'Deployed automotive security modules in partnership with Hyundai AutoEver',
    (21, 0): 'Worked on-site with global team members and Hyundai AutoEver staff, providing technical support to partner companies',
    (22, 0): 'Gamasot Online English | Founder & Operator\t',
    (23, 0): 'Founded and ran a video-call English education business for 4 years; built Unity-based learning games as the sole developer and iterated the product on direct user feedback',
    (26, 0): 'Education',
    (28, 0): 'UNIX systems, network programming, self-directed problem-solving projects',
    (29, 0): 'Pukyong National University',
    (29, 1): ' — B.S. in Industrial Engineering (2007.03 – 2014.02)',
}

DELETE = [24]  # P23에 합쳐서 삭제

DATE_HEADINGS = [10, 16, 19, 22]
RIGHT_TAB = Twips(10740)

doc = docx.Document(SRC)

for (pi, ri), text in R.items():
    doc.paragraphs[pi].runs[ri].text = text

for pi in DATE_HEADINGS:
    ts = doc.paragraphs[pi].paragraph_format.tab_stops
    ts.clear_all()
    ts.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)

for pi in sorted(DELETE, reverse=True):
    el = doc.paragraphs[pi]._element
    el.getparent().remove(el)

doc.save(DST)
print('saved:', DST)

d2 = docx.Document(DST)
for i, p in enumerate(d2.paragraphs):
    if p.text.strip():
        print(f'[P{i}] {p.text}')
