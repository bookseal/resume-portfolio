# -*- coding: utf-8 -*-
"""이기찬_이력서_빌더_업스테이지 (1).docx → 영어 버전 생성.
run 단위 치환으로 서식 보존, 날짜 정렬은 우측 탭스톱으로 재구성."""
import docx
from docx.shared import Twips
from docx.enum.text import WD_TAB_ALIGNMENT

SRC = '이기찬_이력서_빌더_업스테이지 (1).docx'
DST = 'Gichan_Lee_Resume_Builder_Upstage_EN.docx'

# (paragraph_index, run_index) -> new text
R = {
    (0, 0): 'Gichan Lee',
    (1, 0): 'AI Product Builder & Educator - A builder who makes what he builds understandable to everyone',
    (2, 0): '+82 10-6666-2824 | ',
    (2, 1): ' | Jagok-dong, Gangnam-gu, Seoul, South Korea',
    (5, 0): 'Current Lead AI FDE (AI Agent Lead)',
    (5, 1): ": Leading a 3-person team, translating executives' business needs into technology and steering development direction with a DevRel-minded management approach",
    (6, 0): 'Educational Content & Building in Public',
    (6, 1): ': Seoul APT ML 10-step series, Daily Seongsu MLOps guide, built BookToss in public with Upstage Solar + shared product feedback',
    (7, 0): '#1 in Non-Developer Communication & Troubleshooting',
    (7, 1): ': Ranked #1 on the team across all metrics (SRR 92.7%, 0% dissatisfaction) in MS M365 technical support, proving outstanding communication and problem-solving skills',
    (9, 0): 'Experience',
    (10, 0): 'Korea Institute of Business Analysis (KIBA) | Lead AI FDE',
    (10, 1): '\t',
    (10, 2): '2026.05 - Present',
    (11, 0): '3-Role Orchestration:',
    (11, 1): " Designed and operate a collaboration system that connects three roles into a single loop: non-developer clients (intent), AI agents (execution), and verifying developers. Made agent work clear enough for non-developers to drive, and encoded developers' judgment into explicit policies. (Details: 'KIBA Collaboration System' diagram at bit-habit.com)",
    (12, 0): 'Human-in-the-Loop Guardrails: ',
    (12, 1): 'Branched agent actions by type — reads/lookups run automatically, board writes require per-item human approval, and deletions, permission changes, and code changes are prohibited. Set policy to keep meeting notes private and mask PII before commits.',
    (13, 0): 'Claude Code + GitHub Pipeline: ',
    (13, 1): 'Built an asynchronous collaboration flow from meeting notes → AI agent → board (Issues/PRs), which took hold reliably despite the 12-hour US–Korea time difference. Migrated from Plane to GitHub Projects after a build-vs-buy assessment.',
    (14, 0): 'Product from Planning to Deployment: ',
    (14, 1): 'Built \'Quali-fit\', an internal certification-based staffing recommendation system — scoring "who to assign to this task, and why" against certification criteria. Cut a 10-hour task to 2 hours with an LLM workflow for analyzing government guideline documents.',
    (15, 0): 'Technical Leadership: ',
    (15, 1): 'Leading a 3-person team, onboarding juniors to cloud and Git, and running role-play-style mock demos for executives who don\'t read code.',
    (16, 1): '\t',
    (17, 0): 'Achieved a 5.0 average rating, 0% dissatisfaction, and a 92.7% Service Request Resolution (SRR) rate among 20+ engineers (#1 across all metrics)',
    (18, 0): "Founded the internal 'Fireteam' study group, mentoring colleagues to technical certifications (MS-900), and served as a global product tester",
    (19, 1): '\t',
    (19, 2): '2024.09 – 2025.07',
    (20, 0): 'Deployed automotive security modules in partnership with Hyundai AutoEver',
    (21, 0): 'Collaborated with global team members and Hyundai AutoEver staff, providing technical support to partner companies',
    (22, 0): 'Gamasot Online English | Founder & Operator\t',
    (23, 0): 'Founded and ran a video-call English education business for 4 years, mentoring non-native (Filipino) teachers and coaching teaching methods',
    (24, 0): 'As a developer, built Unity-based English learning games for students and fed immediate user feedback back into the curriculum — core DevRel skills learned hands-on',
    (26, 0): 'Education',
    (28, 0): 'UNIX systems, network programming, self-directed problem-solving projects',
    (29, 0): 'Pukyong National University',
    (29, 1): ' — B.S. in Industrial Engineering (2007.03 – 2014.02)',
}

# 날짜 우측 정렬용 탭스톱을 걸 제목 문단들
DATE_HEADINGS = [10, 16, 19, 22]
# A4(11906) - 좌우 여백(약 567*2) = 사용 가능 폭 끝 지점
RIGHT_TAB = Twips(10740)

doc = docx.Document(SRC)

for (pi, ri), text in R.items():
    doc.paragraphs[pi].runs[ri].text = text

for pi in DATE_HEADINGS:
    ts = doc.paragraphs[pi].paragraph_format.tab_stops
    ts.clear_all()
    ts.add_tab_stop(RIGHT_TAB, WD_TAB_ALIGNMENT.RIGHT)

doc.save(DST)
print('saved:', DST)

# 검증 출력
d2 = docx.Document(DST)
for i, p in enumerate(d2.paragraphs):
    if p.text.strip():
        print(f'[P{i}] {p.text}')
